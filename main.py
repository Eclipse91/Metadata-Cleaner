import os
import json
import shutil
import zipfile
import warnings
import subprocess
from datetime import datetime
import mobi
import docx
import base64
import PyPDF2
import pypandoc
from PIL import Image
from mutagen import File as MutagenFile
from pymediainfo import MediaInfo
from openpyxl import load_workbook
from pptx import Presentation
from ebooklib import epub
import xml.etree.ElementTree as ET

## IMAGE

def remove_svg_metadata(svg_path, output_svg_path):
    '''
    Remove metadata from an SVG file.
    '''
    try:
        # Parse the SVG file
        tree = ET.parse(svg_path)
        root = tree.getroot()

        # Find and remove metadata elements
        for metadata in root.findall('{http://www.w3.org/2000/svg}metadata'):
            root.remove(metadata)

        # Write the cleaned SVG content to a new file
        tree.write(output_svg_path)
    except ET.ParseError as e:
        print(f'Failed to parse SVG file: {e}')
    except Exception as e:
        print(f'Failed to remove SVG metadata: {e}')

def remove_image_metadata(image_path, output_image_path):
    '''
    Remove metadata from an image file.
    '''
    try:
        if image_path[-3:].lower() == 'svg':
            remove_svg_metadata(image_path, output_image_path)
        else:
            with Image.open(image_path) as img:
                img_data = list(img.getdata())
                img_without_metadata = Image.new(img.mode, img.size)
                img_without_metadata.putdata(img_data)
                img_without_metadata.save(output_image_path)
    except Exception as e:
        print(f'Failed to remove image metadata: {e}')

def extract_svg_metadata(svg_path):
    '''
    Extract metadata from an SVG file.
    '''
    try:
        tree = ET.parse(svg_path)
        root = tree.getroot()
        
        metadata = {}
        metadata['width'] = root.attrib.get('width', 'unknown')
        metadata['height'] = root.attrib.get('height', 'unknown')
        metadata['viewBox'] = root.attrib.get('viewBox', 'unknown')
        
        # Extract metadata elements if available
        for elem in root.findall('{http://www.w3.org/2000/svg}metadata'):
            metadata.update(elem.attrib)
        
        return metadata
    except ET.ParseError as e:
        print(f'Failed to parse SVG file: {e}')
        return {}
    except Exception as e:
        print(f'Failed to extract SVG metadata: {e}')
        return {}

def extract_gif_metadata(image_path):
    '''
    Extract metadata from a GIF file.
    '''
    try:
        with Image.open(image_path) as img:
            metadata = img.info

        # Convert binary data in metadata to a readable format if necessary
        formatted_metadata = {}
        for key, value in metadata.items():
            if isinstance(value, bytes):
                # Decode bytes to string
                formatted_metadata[key] = value.decode('latin1')  # or 'utf-8', depending on your needs
            elif isinstance(value, tuple):
                # Handle tuple format, especially for 'extension'
                if len(value) == 2 and isinstance(value[0], bytes):
                    formatted_metadata[key] = (value[0].decode('latin1'), value[1])
                else:
                    formatted_metadata[key] = value
            elif isinstance(value, list):
                # Handle list format, assuming it contains binary data
                formatted_metadata[key] = [v.decode('latin1') if isinstance(v, bytes) else v for v in value]
            else:
                formatted_metadata[key] = value

        return formatted_metadata
    except Exception as e:
        print(f'Failed to extract image metadata: {e}')
        return {}

def extract_image_metadata(image_path):
    '''
    Extract metadata from an image file.
    '''
    try:
        if image_path[-3:].lower() == 'svg':
            metadata = extract_svg_metadata(image_path)
        elif image_path[-3:].lower() == 'gif':
            metadata = extract_gif_metadata(image_path)
        else:
            with Image.open(image_path) as img:
                metadata = img.info
        return metadata
    except Exception as e:
        print(f'Failed to extract image metadata: {e}')
        return {}

## AUDIO

def extract_audio_metadata(audio_path):
    '''
    Extract metadata from an audio file.
    '''
    try:
        audio = MutagenFile(audio_path)
        metadata = {k: str(v) for k, v in audio.tags.items()} if audio.tags else {}
        return metadata
    except Exception as e:
        print(f'Failed to extract audio metadata: {e}')
        return {}

def remove_audio_metadata(audio_path):
    '''
    Remove metadata from an audio file.
    '''
    try:
        audio = MutagenFile(audio_path)
        audio.delete()
    except Exception as e:
        print(f'Failed to remove audio metadata: {e}')
    
def copy_file(src_path, dst_path):
    '''
    Copies a file from src_path to dst_path.
    '''
    try:
        shutil.copy(src_path, dst_path)
        # print(f'File copied from {src_path} to {dst_path}')
    except FileNotFoundError:
        print(f'File not found: {src_path}')
    except PermissionError:
        print(f'Permission denied: {dst_path}')
    except Exception as e:
        print(f'Error occurred while copying file: {e}')

def rename_file(current_path, new_path):
    '''
    Rename a file from current_path to new_path.
    '''
    try:
        os.rename(current_path, new_path)
        print(f'File renamed from {current_path} to {new_path}')
    except FileNotFoundError:
        print(f'The file {current_path} does not exist.')
    except PermissionError:
        print(f'Permission denied to rename {current_path}.')
    except Exception as e:
        print(f'An error occurred: {e}')

## VIDEO

def extract_video_metadata(video_path):
    '''
    Extract metadata from a video file.
    '''
    try:
        media_info = MediaInfo.parse(video_path)
        metadata = media_info.to_data()
        return metadata
    except Exception as e:
        print(f'Failed to extract video metadata: {e}')
        return {}

def remove_video_metadata(video_path, output_video_path):
    '''
    Remove metadata from a video file.
    '''
    try:
        # Use ffmpeg to remove metadata
        command = [
            'ffmpeg', '-i', video_path, '-map', '0', '-map_metadata', '-1',
            '-c:v', 'copy', '-c:a', 'copy', output_video_path
        ]
        subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        print(f'Metadata removed and saved to {output_video_path}')
    except subprocess.CalledProcessError as e:
        print(f'Failed to remove video metadata: {e}')

## DOCUMENTS

# def remove_pdf_metadata(pdf_path, output_pdf_path):
#     '''
#     Remove metadata from a PDF file.
#     '''
#     try:
#         reader = PdfReader(pdf_path)
#         writer = PdfWriter()

#         for page in reader.pages:
#             writer.add_page(page)

#         with open(output_pdf_path, 'wb') as f:
#             writer.write(f)
#     except Exception as e:
#         print(f'Failed to remove PDF metadata: {e}')

def extract_pdf_metadata(file_path):
    '''
    Extract metadata from a PDF file.
    '''
    metadata = {}
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            info = reader.getDocumentInfo()
            if info:
                metadata = {key: info[key] for key in info}
    except Exception as e:
        print(f'Failed to extract PDF metadata: {e}')
    return metadata

def extract_docx_metadata(file_path):
    '''
    Extract metadata from a DOCX file.
    '''
    metadata = {}
    try:
        doc = docx.Document(file_path)
        core_props = doc.core_properties
        metadata = {prop: getattr(core_props, prop) for prop in dir(core_props) if not prop.startswith('_')}
    except Exception as e:
        print(f'Failed to extract DOCX metadata: {e}')
    return metadata

# def extract_pdf_metadata(pdf_path):
#     '''
#     Extract metadata from a PDF file.
#     '''
#     try:
#         reader = PdfReader(pdf_path)
#         metadata = reader.metadata
#         return metadata
#     except Exception as e:
#         print(f'Failed to extract PDF metadata: {e}')
#         return {}

def extract_pdf_metadata(file_path):
    '''
    Extract metadata from a PDF file.
    '''
    metadata = {}
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            info = reader.getDocumentInfo()
            if info:
                metadata = {key: info[key] for key in info}
    except Exception as e:
        print(f'Failed to extract PDF metadata: {e}')
    return metadata

def extract_docx_metadata(file_path):
    '''
    Extract metadata from a DOCX file.
    '''
    metadata = {}
    try:
        doc = docx.Document(file_path)
        core_props = doc.core_properties
        metadata = {prop: getattr(core_props, prop) for prop in dir(core_props) if not prop.startswith('_')}
    except Exception as e:
        print(f'Failed to extract DOCX metadata: {e}')
    return metadata

def extract_xlsx_metadata(file_path):
    '''
    Extract metadata from an XLSX file.
    '''
    metadata = {}
    try:
        wb = load_workbook(file_path)
        props = wb.properties
        metadata = {prop: getattr(props, prop) for prop in dir(props) if not prop.startswith('_')}
    except Exception as e:
        print(f'Failed to extract XLSX metadata: {e}')
    return metadata

def extract_pptx_metadata(file_path):
    '''
    Extract metadata from a PPTX file.
    '''
    metadata = {}
    try:
        prs = Presentation(file_path)
        props = prs.core_properties
        metadata = {prop: getattr(props, prop) for prop in dir(props) if not prop.startswith('_')}
    except Exception as e:
        print(f'Failed to extract PPTX metadata: {e}')
    return metadata

def extract_text_metadata(file_path):
    '''
    Extract metadata from a text file.
    '''
    metadata = {}
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        metadata['content'] = content[:500]  # Extracting first 500 characters as sample metadata
    except Exception as e:
        print(f'Failed to extract text metadata: {e}')
    return metadata

def extract_odt_metadata(file_path):
    '''
    Extract metadata from an ODT file.
    '''
    metadata = {}
    try:
        output = pypandoc.convert_file(file_path, 'plain')
        metadata['content'] = output[:500]  # Extracting first 500 characters as sample metadata
    except Exception as e:
        print(f'Failed to extract ODT metadata: {e}')
    return metadata

def extract_rtf_metadata(file_path):
    '''
    Extract metadata from an RTF file.
    '''
    metadata = {}
    try:
        output = pypandoc.convert_file(file_path, 'plain')
        metadata['content'] = output[:500]  # Extracting first 500 characters as sample metadata
    except Exception as e:
        print(f'Failed to extract RTF metadata: {e}')
    return metadata

def extract_html_metadata(file_path):
    '''
    Extract metadata from an HTML file.
    '''
    metadata = {}
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        metadata['content'] = content[:500]  # Extracting first 500 characters as sample metadata
    except Exception as e:
        print(f'Failed to extract HTML metadata: {e}')
    return metadata

def extract_md_metadata(file_path):
    '''
    Extract metadata from a Markdown file.
    '''
    metadata = {}
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        metadata['content'] = content[:500]  # Extracting first 500 characters as sample metadata
    except Exception as e:
        print(f'Failed to extract Markdown metadata: {e}')
    return metadata

def extract_epub_metadata(file_path):
    '''
    Extract metadata from an EPUB file.
    '''
    metadata = {}
    try:
        # Read the EPUB file
        # Suppress specific warning
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", (UserWarning, FutureWarning))
            book = epub.read_epub(file_path)

        # Extract metadata from the EPUB file
        # The `get_metadata` method retrieves metadata from the DC namespace
        title_metadata = book.get_metadata('DC', 'title')
        author_metadata = book.get_metadata('DC', 'creator')
        language_metadata = book.get_metadata('DC', 'language')
        identifier_metadata = book.get_metadata('DC', 'identifier')
        
        metadata['title'] = title_metadata[0][0] if title_metadata else 'Unknown Title'
        metadata['author'] = author_metadata[0][0] if author_metadata else 'Unknown Author'
        metadata['language'] = language_metadata[0][0] if language_metadata else 'Unknown Language'
        metadata['identifier'] = identifier_metadata[0][0] if identifier_metadata else 'Unknown Identifier'

    except Exception as e:
        print(f'Failed to extract EPUB metadata: {e}')
    
    return metadata

def extract_mobi_metadata(file_path):
    '''
    Extract metadata from a MOBI file.
    '''
    metadata = {}
    try:
        book = mobi.read(file_path)
        metadata['title'] = book.title
        metadata['author'] = book.author
    except Exception as e:
        print(f'Failed to extract MOBI metadata: {e}')
    return metadata

def extract_document_metadata(file_path):
    '''
    Extract metadata from a document based on its file extension.
    '''
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.pdf':
        return extract_pdf_metadata(file_path)
    elif file_extension in ['.doc', '.docx']:
        return extract_docx_metadata(file_path)
    elif file_extension in ['.xls', '.xlsx']:
        return extract_xlsx_metadata(file_path)
    elif file_extension in ['.ppt', '.pptx']:
        return extract_pptx_metadata(file_path)
    elif file_extension == '.txt':
        return extract_text_metadata(file_path)
    elif file_extension == '.odt':
        return extract_odt_metadata(file_path)
    elif file_extension == '.rtf':
        return extract_rtf_metadata(file_path)
    elif file_extension == '.html':
        return extract_html_metadata(file_path)
    elif file_extension == '.md':
        return extract_md_metadata(file_path)
    elif file_extension == '.epub':
        return extract_epub_metadata(file_path)
    elif file_extension == '.mobi':
        return extract_mobi_metadata(file_path)
    else:
        return {}

def remove_document_metadata(file_path, output_file_path):
    '''
    Remove metadata from a document.
    '''
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.pdf':
        remove_pdf_metadata(file_path, output_file_path)
    elif file_extension in ['.doc', '.docx']:
        remove_docx_metadata(file_path, output_file_path)
    elif file_extension in ['.xls', '.xlsx']:
        remove_xlsx_metadata(file_path, output_file_path)
    elif file_extension in ['.ppt', '.pptx']:
        remove_pptx_metadata(file_path, output_file_path)
    elif file_extension in ['.epub']:
        remove_epub_metadata(file_path, output_file_path)
    else:
        print(f'Unsupported document extension: {file_extension}')

def remove_pdf_metadata(file_path, output_file_path):
    '''
    Remove metadata from a PDF file.
    '''
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            writer = PyPDF2.PdfFileWriter()
            for i in range(reader.getNumPages()):
                writer.addPage(reader.getPage(i))
            writer.addMetadata({})
            with open(output_file_path, 'wb') as out_f:
                writer.write(out_f)
        print(f'Metadata removed from PDF: {output_file_path}')
    except Exception as e:
        print(f'Failed to remove PDF metadata: {e}')

def remove_docx_metadata(file_path, output_file_path):
    '''
    Remove metadata from a DOCX file.
    '''
    try:
        doc = docx.Document(file_path)
        core_props = doc.core_properties
        for prop in dir(core_props):
            if not prop.startswith('_') and not callable(getattr(core_props, prop)):
                setattr(core_props, prop, None)
        doc.save(output_file_path)
        print(f'Metadata removed from DOCX: {output_file_path}')
    except Exception as e:
        print(f'Failed to remove DOCX metadata: {e}')

def remove_pptx_metadata(file_path, output_file_path):
    '''
    Remove metadata from a PPTX file.
    '''
    try:
        prs = Presentation(file_path)
        props = prs.core_properties
        for prop in dir(props):
            if not prop.startswith('_') and not callable(getattr(props, prop)):
                setattr(props, prop, None)
        prs.save(output_file_path)
        print(f'Metadata removed from PPTX: {output_file_path}')
    except Exception as e:
        print(f'Failed to remove PPTX metadata: {e}')

def remove_xlsx_metadata(file_path, output_file_path):
    '''
    Remove metadata from an XLSX file.
    '''
    try:
        wb = load_workbook(file_path)
        wb.properties = None
        wb.save(output_file_path)
        print(f'Metadata removed from XLSX: {output_file_path}')
    except Exception as e:
        print(f'Failed to remove XLSX metadata: {e}')

def remove_epub_metadata(file_path, output_file_path):
    '''
    Remove metadata from an EPUB file.
    '''
    try:
        # Create a temporary directory to work with the EPUB file
        temp_dir = 'temp_epub'
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir, exist_ok=True)
        
        # Extract the EPUB file
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Load the content.opf file
        opf_path = os.path.join(temp_dir, 'OEBPS', 'content.opf')
        if not os.path.isfile(opf_path):
            opf_path = os.path.join(temp_dir, 'content.opf')  # Handle cases where it's directly in root
        
        # Parse and modify the content.opf file
        tree = ET.parse(opf_path)
        root = tree.getroot()
        ns = {'dc': 'http://purl.org/dc/elements/1.1/'}
        
        # Collect elements to remove
        elements_to_remove = root.findall('.//dc:*', ns)
        
        for elem in elements_to_remove:
            parent = elem.find('..')  # Find the parent node of the element
            if parent is not None:
                parent.remove(elem)
        
        # Save the modified content.opf file
        tree.write(opf_path, encoding='utf-8', xml_declaration=True)
        
        # Create a new EPUB file
        with zipfile.ZipFile(output_file_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for foldername, subfolders, filenames in os.walk(temp_dir):
                for filename in filenames:
                    filepath = os.path.join(foldername, filename)
                    arcname = os.path.relpath(filepath, temp_dir)
                    zip_ref.write(filepath, arcname)
        
        # Clean up temporary directory
        shutil.rmtree(temp_dir)
            
    except Exception as e:
        print(f'Failed to remove EPUB metadata: {e}')

## MAIN

def save_metadata_to_file(metadata, metadata_file_path):
    '''
    Save metadata to a file.
    '''
    try:
        # Encode the bytes objects in the metadata dictionary
        metadata = encode_bytes_in_dict(metadata)
        
        # Save the metadata to a file as JSON
        with open(metadata_file_path, 'w') as f:
            json.dump(metadata, f, indent=4)
    except Exception as e:
        print(f'Failed to save metadata to file: {e}')

def encode_bytes_in_dict(obj):
    '''
    Encode bytes objects in a dictionary to base64 strings.
    '''
    if isinstance(obj, dict):
        return {k: encode_bytes_in_dict(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [encode_bytes_in_dict(v) for v in obj]
    elif isinstance(obj, bytes):
        return base64.b64encode(obj).decode('utf-8')
    else:
        return obj

def results_configurator(file_name):
    '''
    Configure the folder where to put all the resulting files.
    '''
    results_directory = './results/'
    current_datetime = datetime.now()
    formatted_datetime = current_datetime.strftime('%Y%m%d_%H%M%S')
    folder_name = results_directory + formatted_datetime + '_' + file_name.replace('.','_')
    os.makedirs(folder_name, exist_ok=True)

    return folder_name
    
def list_files_in_current_folder():
    '''
    Lists all files in the current directory.
    '''
    try:
        files = [f for f in os.listdir('.') if os.path.isfile(f)]
        return files
    except Exception as e:
        print(f'Error occurred while listing files: {e}')
        return []

def process_file(file_path, folder_name):
    '''
    Analyze the file and remove any metadata.
    '''
    file_extension = os.path.splitext(file_path)[1].lower()
    new_file_path = folder_name + '/' + file_path
    file_path = './clean/' + file_path
    metadata_file_path = f'{new_file_path}_metadata.json'
    output_file_path = f'{new_file_path}_no_metadata{file_extension}'
    
    try:
        if file_extension.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg']: # , '.tiff', '.heic', '.raw']:
            metadata = extract_image_metadata(file_path)
            save_metadata_to_file(metadata, metadata_file_path)
            remove_image_metadata(file_path, output_file_path)
        elif file_extension.lower() in ['.mp3', '.wav', '.ogg']: # , '.m4a', '.aac', '.wma', '.alac', '.aiff', '.flac']:
            copy_file(file_path, folder_name)
            rename_file(new_file_path, output_file_path)
            metadata = extract_audio_metadata(output_file_path)
            save_metadata_to_file(metadata, metadata_file_path)
            remove_audio_metadata(output_file_path)
        elif file_extension.lower() in ['.mp4', '.mkv', '.avi', '.mov', '.webm']: #, '.mpeg', '.mpg', '.3gp', '.wmv', '.flv']:
            metadata = extract_video_metadata(file_path)
            save_metadata_to_file(metadata, metadata_file_path)
            remove_video_metadata(file_path, output_file_path)
        # elif file_extension.lower() in ['.pdf', '.epub', '.mobi', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.txt', '.odt', '.rtf', '.html', '.md']:
        #     metadata = extract_document_metadata(file_path)
        #     save_metadata_to_file(metadata, metadata_file_path)
        #     remove_document_metadata(file_path, output_file_path)
        else:
            print(f'Unsupported file type: {file_extension}')
            return
        
        print(f'Metadata saved to: {metadata_file_path}')
        print(f'File without metadata saved to: {output_file_path}\n')
    except Exception as e:
        print(f'An error occurred while processing the file: {e}')

def execution_time(func):
    '''
    Decorator that prints the current date and time before and after
    executing the given function, and measures the time taken for execution.
    The datetime format is 'YYYYMMDD_HHMMSS'.
    '''
    def wrapper():
        current_datetime = datetime.now()
        formatted_datetime = current_datetime.strftime('%Y%m%d_%H%M%S')
        print(f'Program started at {formatted_datetime}')
        func()
        current_datetime = datetime.now()
        formatted_datetime = current_datetime.strftime('%Y%m%d_%H%M%S')
        print(f'Program ended at {formatted_datetime}')
    return wrapper

@execution_time
def main():
    files = [f for f in os.listdir('./clean') if os.path.isfile(os.path.join('./clean', f))]
    for file in files:
        folder_name = results_configurator(file)

        try:
            process_file(file, folder_name)
        except Exception as e:
            print(f'An error occurred: {e}')

if __name__ == '__main__':
    main()

